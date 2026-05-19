#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
make_word.py

make_json.py が作成した work/{subject}.json から、学校提出用の Word(docx) を作成する。

実行例:
    python scripts/make_word.py 2031002

出力例:
    exam/word/A/2031002_A.docx
    exam/word/B/2031002_B.docx

方針:
    - PDFと完全一致は目指さない
    - Wordで編集しやすい形を優先する
    - LaTeX数式は第1段階では文字列として残す
    - 画像は可能なら貼り付ける
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from exam_utils import add_subject_arg, load_exam_context


# ------------------------------------------------------------
# 基本設定
# ------------------------------------------------------------
DEFAULT_FONT = "游明朝"
DEFAULT_CODE_FONT = "Consolas"


# ------------------------------------------------------------
# Wordスタイル補助
# ------------------------------------------------------------
def set_run_font(run, font_name: str = DEFAULT_FONT, size_pt: int | None = None) -> None:
    """日本語フォントを含めて run にフォント設定する。"""
    run.font.name = font_name
    if size_pt is not None:
        run.font.size = Pt(size_pt)

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_font(paragraph, font_name: str = DEFAULT_FONT, size_pt: int | None = None) -> None:
    for run in paragraph.runs:
        set_run_font(run, font_name, size_pt)


def set_document_defaults(doc: Document) -> None:
    """Word文書全体の基本スタイルを設定する。"""
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = DEFAULT_FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), DEFAULT_FONT)

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)


def add_heading_text(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(level=level)
    p.text = text
    set_paragraph_font(p, DEFAULT_FONT, 14 if level == 1 else 12)


def add_normal_paragraph(doc: Document, text: str = ""):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, DEFAULT_FONT, 10)
    return p


def add_code_block(doc: Document, lines: list[str]) -> None:
    """コードブロックを等幅フォントで出す。"""
    if not lines:
        return

    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    text = "\n".join(str(x) for x in lines)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, DEFAULT_CODE_FONT, 9)


# ------------------------------------------------------------
# テキスト整形
# ------------------------------------------------------------
def clean_text(text: Any) -> str:
    """
    JSON内の文字列をWord用に軽く整える。
    第1段階ではLaTeX数式は文字列として残す。
    """
    if text is None:
        return ""
    s = str(text)

    # raw LaTeX [[...]] は中身だけにする
    s = re.sub(r"\[\[(.*?)\]\]", r"\1", s, flags=re.DOTALL)

    # Wordで編集しやすいように、過度な空白だけ調整
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    return s


def add_text_with_linebreaks(doc: Document, text: Any) -> None:
    s = clean_text(text)
    lines = s.split("\n")
    if not lines:
        add_normal_paragraph(doc, "")
        return
    p = doc.add_paragraph()
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, DEFAULT_FONT, 10)


# ------------------------------------------------------------
# JSON読み込み・検証
# ------------------------------------------------------------
def load_json(json_path: Path) -> dict[str, Any]:
    if not json_path.exists():
        raise FileNotFoundError(
            "JSONファイルが見つかりません。\n"
            "先に make_json.py を実行してください。\n"
            f"JSON path: {json_path}"
        )
    with json_path.open("r", encoding="utf-8") as f:
        return json.load(f)


def get_versions(data: dict[str, Any]) -> list[dict[str, Any]]:
    versions = data.get("versions") or []
    if not versions:
        raise RuntimeError("JSON内に versions がありません。")
    return versions


# ------------------------------------------------------------
# 画像処理
# ------------------------------------------------------------
def find_image_path(image_name: str, exam_dir: Path) -> Path | None:
    """画像ファイルを探す。"""
    if not image_name:
        return None

    candidates = [
        Path(image_name),
        exam_dir / "images" / image_name,
        exam_dir.parent / "images" / image_name,
        exam_dir / image_name,
    ]

    for path in candidates:
        if path.exists():
            return path
    return None


def add_image(doc: Document, item: dict[str, Any], exam_dir: Path) -> None:
    path_value = item.get("path") or ""
    image_path = find_image_path(str(path_value), exam_dir)

    if image_path is None:
        add_normal_paragraph(doc, f"[画像が見つかりません: {path_value}]")
        return

    width_ratio = item.get("width", 0.8)
    try:
        width_ratio = float(width_ratio)
    except Exception:
        width_ratio = 0.8

    # A4本文幅をざっくり6.5インチとして計算
    width_inches = max(1.0, min(6.5, 6.5 * width_ratio))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture(str(image_path), width=Inches(width_inches))
    except Exception as e:
        add_normal_paragraph(doc, f"[画像挿入エラー: {image_path} / {e}]")


# ------------------------------------------------------------
# 要素レンダリング
# ------------------------------------------------------------
def render_cover(doc: Document, item: dict[str, Any], version: str) -> None:
    title = clean_text(item.get("title") or item.get("subject") or "試験問題")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    set_run_font(run, DEFAULT_FONT, 18)

    subject = item.get("subject")
    fsyear = item.get("fsyear")
    if subject or fsyear:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(" / ".join(str(x) for x in [fsyear, subject] if x))
        set_run_font(run, DEFAULT_FONT, 11)

    if version:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Version {version}")
        set_run_font(run, DEFAULT_FONT, 11)

    notes = item.get("notes") or []
    for note in notes:
        add_text_with_linebreaks(doc, note)

    doc.add_page_break()


def render_premise(doc: Document, item: dict[str, Any], exam_dir: Path) -> None:
    title = item.get("title") or "前提条件"
    p = doc.add_paragraph()
    run = p.add_run(f"【{title}】")
    run.bold = True
    set_run_font(run, DEFAULT_FONT, 11)

    for sub in item.get("content") or []:
        render_content_item(doc, sub, exam_dir)

    add_normal_paragraph(doc, "")


def render_choices(doc: Document, item: dict[str, Any]) -> None:
    values = item.get("values") or []
    if not values:
        return

    for idx, v in enumerate(values, start=1):
        if isinstance(v, dict):
            label = v.get("label") or str(idx)
            text = v.get("text") or ""
        else:
            label = str(idx)
            text = str(v)
        add_normal_paragraph(doc, f"{label}. {clean_text(text)}")


def render_content_item(doc: Document, item: Any, exam_dir: Path) -> None:
    if not isinstance(item, dict):
        add_text_with_linebreaks(doc, item)
        return

    typ = item.get("type")

    if typ in ("text", "sline"):
        add_text_with_linebreaks(doc, item.get("value") or "")

    elif typ == "multiline":
        for v in item.get("values") or []:
            add_text_with_linebreaks(doc, v)

    elif typ == "preline":
        for v in item.get("values") or []:
            add_text_with_linebreaks(doc, v)

    elif typ == "code":
        add_code_block(doc, item.get("lines") or [])

    elif typ == "choices":
        render_choices(doc, item)

    elif typ == "image":
        add_image(doc, item, exam_dir)

    elif typ == "vspace":
        add_normal_paragraph(doc, "")

    elif typ == "pagebreak":
        #doc.add_page_break()
        return

    elif typ == "premise":
        render_premise(doc, item, exam_dir)

    elif typ == "cover":
        # coverはversion単位の先頭で処理するため、ここでは何もしない
        return

    elif typ == "metainfo":
        return

    else:
        add_normal_paragraph(doc, f"[未対応type: {typ}] {item}")


def render_question(doc: Document, q: dict[str, Any], exam_dir: Path) -> None:
    number = q.get("number") or "?"
    qtext = clean_text(q.get("question") or "")

    p = doc.add_paragraph()
    run = p.add_run(f"問{number}：{qtext}")
    run.bold = True
    set_run_font(run, DEFAULT_FONT, 11)

    for item in q.get("content") or []:
        render_content_item(doc, item, exam_dir)

    for sub in q.get("subquestions") or []:
        render_subquestion(doc, sub, number, exam_dir)

    add_normal_paragraph(doc, "")


def render_subquestion(doc: Document, sub: dict[str, Any], parent_number: str, exam_dir: Path) -> None:
    number = sub.get("number") or "?"
    qtext = clean_text(sub.get("question") or "")

    p = doc.add_paragraph()
    run = p.add_run(f"問{parent_number}-{number}：{qtext}")
    run.bold = True
    set_run_font(run, DEFAULT_FONT, 10)

    for item in sub.get("content") or []:
        render_content_item(doc, item, exam_dir)


# ------------------------------------------------------------
# Word生成
# ------------------------------------------------------------
def build_docx_for_version(
    *,
    subject: str,
    version_block: dict[str, Any],
    exam_dir: Path,
    out_path: Path,
) -> None:
    version = str(version_block.get("version") or "A")
    questions = version_block.get("questions") or []

    doc = Document()
    set_document_defaults(doc)

    # coverは常に表示する方針
    cover_items = [q for q in questions if isinstance(q, dict) and q.get("type") == "cover"]
    if cover_items:
        render_cover(doc, cover_items[0], version)
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{subject} 試験問題（{version}）")
        run.bold = True
        set_run_font(run, DEFAULT_FONT, 18)
        doc.add_page_break()

    for item in questions:
        if not isinstance(item, dict):
            continue

        typ = item.get("type")
        if typ == "cover":
            continue
        elif typ == "premise":
            render_premise(doc, item, exam_dir)
#        elif typ in ("vspace", "pagebreak", "text", "multiline", "code", "choices", "image"):
        elif typ in ("vspace", "text", "multiline", "code", "choices", "image"):
            render_content_item(doc, item, exam_dir)
        elif "question" in item or "subquestions" in item:
            render_question(doc, item, exam_dir)
        else:
            render_content_item(doc, item, exam_dir)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


# ------------------------------------------------------------
# main
# ------------------------------------------------------------
def main() -> None:
    parser = argparse.ArgumentParser(description="JSONから学校提出用Word(docx)を作成します。")
    add_subject_arg(parser)
    args = parser.parse_args()

    exam_context = load_exam_context(args.subject, load_workbook=False)

    subject = exam_context.subject
    json_path = exam_context.work_dir / f"{subject}.json"
    word_dir = exam_context.exam_dir / "wordexcel"

    print(f"科目番号: {subject}")
    print(f"年度: {exam_context.fsyear}")
    print(f"入力JSON: {json_path}")
    print(f"出力Word: {word_dir}")

    data = load_json(json_path)
    versions = get_versions(data)

    for block in versions:
        version = str(block.get("version") or "A")
        out_path = word_dir / f"{subject}_{version}.docx"

        build_docx_for_version(
            subject=subject,
            version_block=block,
            exam_dir=exam_context.exam_dir,
            out_path=out_path,
        )

        print(f"✅ Word作成: {out_path}")

    print("🎯 Word作成が完了しました。")


if __name__ == "__main__":
    try:
        main()
    except SystemExit:
        raise
    except Exception as e:
        if "--debug" in sys.argv:
            import traceback
            traceback.print_exc()
        else:
            print()
            print("🙅🏻‍♂️ エラー:")
            print(e)
        raise SystemExit(1)
